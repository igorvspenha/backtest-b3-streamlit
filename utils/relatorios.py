
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

def gerar_relatorio_pdf(ticker, df, metrics, fig, params, output_path="relatorio_backtest.pdf"):
    img_path = "grafico_temp.png"
    fig.savefig(img_path, dpi=150, bbox_inches="tight")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, f"Relatório de Backtest – {ticker}", ln=True)

    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, f"Estratégia: {params.get('Estratégia', 'N/A')}", ln=True)
    for k, v in params.items():
        if k != "Estratégia":
            pdf.cell(0, 8, f"{k}: {v}", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Métricas de Desempenho", ln=True)
    pdf.set_font("Arial", size=12)
    for k, v in metrics.items():
        pdf.cell(0, 8, f"{k}: {v}", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Resumo da Análise", ln=True)
    pdf.set_font("Arial", size=12)
    resumo = f"O backtest da estratégia no ativo {ticker} entre {df.index.min().date()} e {df.index.max().date()} "
    resumo += f"mostrou um retorno acumulado de {metrics['Retorno Acumulado (%)']}%, com "
    resumo += f"um drawdown máximo de {metrics['Drawdown Máximo (%)']}% e um índice de Sharpe de {metrics['Índice de Sharpe']}."
    pdf.multi_cell(0, 8, resumo)

    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Gráfico com Sinais", ln=True)
    pdf.image(img_path, w=180)

    pdf.output(output_path)
    os.remove(img_path)
    return output_path

def gerar_relatorio_docx(ticker, df, metrics, fig, params, output_path="relatorio_backtest.docx"):
    doc = Document()
    doc.add_heading(f'Relatório de Backtest – {ticker}', 0)

    doc.add_heading('📌 Estratégia e Parâmetros', level=1)
    for k, v in params.items():
        doc.add_paragraph(f'{k}: {v}')

    doc.add_heading('📈 Métricas de Desempenho', level=1)
    for k, v in metrics.items():
        doc.add_paragraph(f'{k}: {v}')

    doc.add_heading('📝 Análise', level=1)
    resumo = f"O ativo {ticker} teve retorno de {metrics['Retorno Acumulado (%)']}%, "
    resumo += f"drawdown máximo de {metrics['Drawdown Máximo (%)']}%, e Sharpe {metrics['Índice de Sharpe']}."
    doc.add_paragraph(resumo)

    img_path = "grafico_docx_temp.png"
    fig.savefig(img_path, dpi=150, bbox_inches='tight')
    doc.add_picture(img_path, width=Inches(6))
    os.remove(img_path)

    doc.save(output_path)
    return output_path
